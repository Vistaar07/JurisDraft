import os
from tqdm import tqdm
import fitz  # PyMuPDF
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
# NEW: OCR imports
import pytesseract
from PIL import Image

# Configure Tesseract path for Windows/WSL/Linux so pytesseract can find the binary
import platform
import shutil
import pathlib
import os as _os

def _configure_tesseract_path():
    # Respect explicit env override if provided
    cmd_override = _os.environ.get("TESSERACT_CMD")
    if cmd_override and pathlib.Path(cmd_override).exists():
        pytesseract.pytesseract.tesseract_cmd = cmd_override
        return

    system = platform.system().lower()

    # Detect WSL specifically
    is_wsl = False
    try:
        with open("/proc/sys/kernel/osrelease", "r") as f:
            is_wsl = "microsoft" in f.read().lower()
    except Exception:
        pass

    if system == "windows":
        win_path = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
        if pathlib.Path(win_path).exists():
            pytesseract.pytesseract.tesseract_cmd = win_path
            # Ensure tessdata is discoverable (helps when using extra language packs)
            _os.environ.setdefault("TESSDATA_PREFIX", r"C:\\Program Files\\Tesseract-OCR")
    elif is_wsl:
        # Use Windows Tesseract from WSL if available (WSL interop)
        wsl_win_path = "/mnt/c/Program Files/Tesseract-OCR/tesseract.exe"
        if pathlib.Path(wsl_win_path).exists():
            pytesseract.pytesseract.tesseract_cmd = wsl_win_path
            _os.environ.setdefault("TESSDATA_PREFIX", "/mnt/c/Program Files/Tesseract-OCR")
        else:
            # Fallback to Linux tesseract installed inside WSL
            cmd = shutil.which("tesseract")
            if cmd:
                pytesseract.pytesseract.tesseract_cmd = cmd
    else:
        # Native Linux/macOS: rely on PATH if present
        cmd = shutil.which("tesseract")
        if cmd:
            pytesseract.pytesseract.tesseract_cmd = cmd

# Run configuration at import time
_configure_tesseract_path()


def ocr_pdf(file_path: str, dpi_scale: float = 2.0, lang: str = "eng") -> str:
    """Render each PDF page to an image and run Tesseract OCR. Returns concatenated text."""
    text_parts = []
    doc = fitz.open(file_path)
    try:
        # Zoom for better OCR accuracy (72 DPI * scale).
        mat = fitz.Matrix(dpi_scale, dpi_scale)
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)  # render page to raster
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            # You can tweak Tesseract page segmentation mode (psm) if layout is multi-column
            # common choices: 3 (auto), 4 (column), 6 (block of text)
            config = "--psm 4"
            page_text = pytesseract.image_to_string(img, lang=lang, config=config)
            if page_text:
                text_parts.append(page_text)
    finally:
        doc.close()
    return "\n".join(text_parts).strip()


def load_and_read_doc(file_path):
    """Loads text from PDF or DOCX file with detailed logging. Falls back to OCR for image PDFs."""
    try:
        if file_path.lower().endswith(".pdf"):
            doc = fitz.open(file_path)
            try:
                # First try native text extraction (fast for digitally generated PDFs)
                text = "".join(page.get_text() for page in doc)
            finally:
                doc.close()

            # Fallback to OCR if no extractable text
            if not text or text.isspace():
                print(f"INFO: No extractable text in {file_path}. Running OCR...")
                # Adjust language codes as needed: e.g., "eng+hin"
                text = ocr_pdf(file_path, dpi_scale=2.0, lang="eng")

        elif file_path.lower().endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
        else:
            print(f"Skipping unsupported file type: {file_path}")
            return None

        # Final sanity check
        if not text or text.isspace():
            print(f"WARNING: OCR/text extraction produced no text for {file_path}.")
            return None

        return text
    except Exception as e:
        print(f"ERROR reading {file_path}: {e}")
        return None


def process_documents_from_path(corpus_path: str) -> list[LangchainDocument]:
    """Processes all PDF/DOCX files in a directory, chunks them, and returns LangChain Documents."""
    all_documents = []
    print(f"Processing documents from: {corpus_path}")

    file_list = []
    for root, _, files in os.walk(corpus_path):
        for file in files:
            file_list.append(os.path.join(root, file))

    for file_path in tqdm(file_list, desc=f"Reading files in {corpus_path}"):
        text = load_and_read_doc(file_path)

        if text:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)

            for i, chunk in enumerate(chunks):
                doc = LangchainDocument(
                    page_content=chunk,
                    metadata={"source": file_path, "chunk_number": i},
                )
                all_documents.append(doc)

    print(f"Processed {len(all_documents)} document chunks.")
    return all_documents


if __name__ == "__main__":
    print("Testing processing for Acts...")
    acts_docs = process_documents_from_path("../data_corpus/Acts")
    print(f"Found {len(acts_docs)} chunks for Acts.")

    print("\nTesting processing for Judgments...")
    judgment_docs = process_documents_from_path("../data_corpus/Judgments")
    print(f"Found {len(judgment_docs)} chunks for Judgments.")