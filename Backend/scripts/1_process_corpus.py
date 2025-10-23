import os
from tqdm import tqdm
import fitz  # PyMuPDF
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema.document import Document as LangchainDocument

def load_and_read_doc(file_path):
    """Loads text from PDF or DOCX file."""
    try:
        if file_path.endswith(".pdf"):
            doc = fitz.open(file_path)
            text = "".join(page.get_text() for page in doc)
            doc.close()
            return text
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            print(f"Skipping unsupported file type: {file_path}")
            return None
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def process_documents_from_path(corpus_path: str) -> list[LangchainDocument]:
    """
    Processes all PDF/DOCX files in a directory, chunks them,
    and returns a list of Langchain Documents.
    """
    all_documents = []
    print(f"Processing documents from: {corpus_path}")

    # Walk through the directory
    for root, _, files in os.walk(corpus_path):
        for file in tqdm(files, desc=f"Reading files in {root}"):
            file_path = os.path.join(root, file)
            text = load_and_read_doc(file_path)

            if text:
                # Use RecursiveCharacterTextSplitter for smart chunking
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len
                )
                chunks = text_splitter.split_text(text)

                # Create documents with metadata
                for i, chunk in enumerate(chunks):
                    doc = LangchainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "chunk_number": i
                        }
                    )
                    all_documents.append(doc)

    print(f"Processed {len(all_documents)} document chunks.")
    return all_documents

if __name__ == "__main__":
    # This allows you to test the script directly
    print("Testing processing for Acts...")
    acts_docs = process_documents_from_path("../data_corpus/Acts")
    print(f"Found {len(acts_docs)} chunks for Acts.")

    print("\nTesting processing for Judgments...")
    judgment_docs = process_documents_from_path("../data_corpus/Judgments")
    print(f"Found {len(judgment_docs)} chunks for Judgments.")